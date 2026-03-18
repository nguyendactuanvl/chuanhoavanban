import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- 1. CẤU HÌNH GIAO DIỆN ---
st.set_page_config(page_title="Chuẩn hóa văn bản hành chính", layout="wide")

# Hàm tạo file Word nâng cao (Xử lý định dạng chuyên nghiệp hơn)
def create_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    
    for line in text.split('\n'):
        p = doc.add_paragraph(line)
        # Nếu là Quốc hiệu hoặc Tên cơ quan (giả định dựa trên nội dung), có thể căn giữa
        if "CỘNG HÒA XÃ HỘI" in line or "Độc lập - Tự do" in line or "TỜ TRÌNH" in line:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

st.markdown("<h1 style='text-align: center;'>📄 Chuẩn Hóa Văn Bản Hành Chính AI</h1>", unsafe_allow_html=True)

with st.sidebar:
    st.title("⚙️ Cấu hình")
    user_api_key = st.text_input("Dán Gemini API Key:", type="password")

if not user_api_key:
    st.warning("👈 Vui lòng nhập API Key!")
    st.stop()

# --- 2. CHỈ DẪN HỆ THỐNG CỰC KỲ CHI TIẾT ---
@st.cache_resource
def get_model(api_key):
    genai.configure(api_key=api_key)
    
    # BÍ QUYẾT Ở ĐÂY: Chỉ dẫn AI cách dàn trang theo hàng và cột
    instruct = """
    `Bạn là một chuyên gia ngôn ngữ học và chuyên gia văn thư lưu trữ tại Việt Nam.
Nhiệm vụ của bạn là:
1. Sửa lỗi chính tả, ngữ pháp, dấu câu và cải thiện văn phong cho đoạn văn bản gốc.
2. Chuẩn hóa đoạn văn bản theo đúng quy định về thể thức văn bản hành chính của Ban Quản lý dự án Đầu tư xây dựng khu vực 1 thành phố Huế và cập nhật theo Nghị định 187/2025/NĐ-CP.

YÊU CẦU VỀ THỂ THỨC (BẮT BUỘC ÁP DỤNG THEO NGHỊ ĐỊNH 187/2025/NĐ-CP):
- Trình bày kết quả dưới dạng HTML để đảm bảo định dạng chính xác khi copy sang Google Docs/Word.
- Font chữ: Times New Roman.
- Cỡ chữ nội dung chính: 13pt hoặc 14pt (thống nhất dùng 14pt).
- Canh lề nội dung chính: Canh đều hai bên (justify), lùi đầu dòng (text-indent) 1cm đến 1.27cm (dùng 1.27cm), khoảng cách dòng (line-height) 1.5, khoảng cách đoạn tối thiểu 6pt. Bắt buộc sử dụng thẻ <div style="text-align: justify; text-indent: 1.27cm; margin-top: 6pt; margin-bottom: 6pt; line-height: 1.5;"> cho mỗi đoạn văn nội dung.
- Bắt buộc dùng bảng HTML không viền cho phần đầu và phần cuối văn bản.
- Từ "Điều", số thứ tự và tên của điều được trình bày bằng chữ in thường, cỡ chữ 14, kiểu chữ đứng, đậm. Ví dụ: <b>Điều 1. Tên điều</b>
- Số thứ tự các khoản trong mỗi điều dùng số Ả Rập, sau số thứ tự có dấu chấm (.), cỡ chữ 14, kiểu chữ đứng.

MẪU PHẦN ĐẦU VĂN BẢN:
<table style="width: 100%; border: none; font-family: 'Times New Roman', serif;">
  <tr>
    <td style="width: 40%; text-align: center; vertical-align: top; font-size: 13pt;">
      UBND THÀNH PHỐ HUẾ<br>
      <b>BAN QUẢN LÝ DỰ ÁN<br>ĐẦU TƯ XÂY DỰNG KHU VỰC 1</b><br>
      <hr style="width: 40%; border-top: 1px solid black; margin: 5px auto;">
      Số: .../QĐ-BQLKV1 <!-- Hoặc ký hiệu tương ứng: /CV-BQLKV1, /TTr-BQLKV1... -->
    </td>
    <td style="width: 60%; text-align: center; vertical-align: top;">
      <b style="font-size: 13pt;">CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</b><br>
      <b style="font-size: 14pt;">Độc lập - Tự do - Hạnh phúc</b><br>
      <hr style="width: 50%; border-top: 1px solid black; margin: 5px auto;">
      <i style="font-size: 14pt;">Huế, ngày ... tháng ... năm ...</i>
    </td>
  </tr>
</table>

MẪU PHẦN CUỐI VĂN BẢN (Nơi nhận và Chữ ký):
<table style="width: 100%; border: none; font-family: 'Times New Roman', serif;">
  <tr>
    <td style="width: 50%; vertical-align: top;">
      <b style="font-size: 12pt; font-style: italic;">Nơi nhận:</b><br>
      <div style="font-size: 11pt; line-height: 1.2;">
      - Như trên;<br>
      - Lưu: VT, TCHC.
      </div>
    </td>
    <td style="width: 50%; text-align: center; vertical-align: top; font-size: 14pt;">
      <b>GIÁM ĐỐC</b><br>
      <i>(Chữ ký, dấu)</i><br><br><br><br>
      <b>Nguyễn Trần Nhật Tuấn</b>
    </td>
  </tr>
</table>

Chỉ trả về mã HTML đã được chuẩn hóa, không giải thích thêm. 
QUAN TRỌNG: 
1. KHÔNG bọc kết quả trong block code markdown (ví dụ: \`\`\`html ... \`\`\`), chỉ trả về mã HTML thuần túy.
2. KHÔNG trả về toàn bộ trang HTML (tuyệt đối không dùng thẻ <!DOCTYPE html>, <html>, <head>, <body>). Chỉ trả về các thẻ HTML chứa nội dung văn bản (như <div>, <table>, <p>, <b>, <i>...).
    """
    
    available = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    selected = 'models/gemini-1.5-flash' if 'models/gemini-1.5-flash' in available else available[0]
    return genai.GenerativeModel(model_name=selected, system_instruction=instruct)

model_ai = get_model(user_api_key)

# --- 3. XỬ LÝ ---
col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("Tải file .docx", type=["docx"])
    user_text = st.text_area("Dán nội dung:", height=400)
    process_btn = st.button("🚀 CHUẨN HÓA NGAY", type="primary", use_container_width=True)

with col2:
    if process_btn:
        input_data = ""
        if uploaded_file:
            doc = Document(uploaded_file)
            input_data = "\n".join([p.text for p in doc.paragraphs])
        else:
            input_data = user_text

        if input_data:
            with st.spinner("Đang chuẩn hóa..."):
                response = model_ai.generate_content(input_data, generation_config={"temperature": 0.1})
                st.session_state['res'] = response.text
                st.markdown(f"```\n{response.text}\n```") # Hiển thị dạng code để dễ copy đúng khoảng cách
                
                st.download_button("📥 Tải xuống Word", create_docx(response.text), "van_ban.docx", use_container_width=True)
